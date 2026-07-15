# -*- coding: utf-8 -*-
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DATA='/mnt/data/KT&KDQT_ DS SV, HV, NCS tốt nghiệp 2025-2026 (03.07).xlsx'
OUT='/mnt/data/Bao_cao_hoc_thuat_Tot_nghiep_2025_2026_visual.docx'
CHART_DIR='/mnt/data/charts_graduation'
os.makedirs(CHART_DIR, exist_ok=True)

# ---------- Data ----------
cn = pd.read_excel(DATA, sheet_name='CN', header=0)
ths = pd.read_excel(DATA, sheet_name='THS', header=1)
ts = pd.read_excel(DATA, sheet_name='TS', header=1)
cn['GPA'] = pd.to_numeric(cn['GPA'], errors='coerce')
cn = cn.dropna(subset=['GPA']).copy()
cn['Loại tốt nghiệp'] = cn['Loại tốt nghiệp'].astype(str).str.strip()
cn['OnTime'] = np.where(cn['Khóa'].astype(str).str.contains('2022'), 'Đúng hạn', 'Trễ hạn')

# ---------- Charts ----------
plt.rcParams.update({'font.size': 10, 'font.family': 'DejaVu Sans'})

def savefig(name):
    path=os.path.join(CHART_DIR,name)
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches='tight')
    plt.close()
    return path

# 1 GPA histogram overall + gender density-style histograms
plt.figure(figsize=(7.2,4.1))
plt.hist(cn.loc[cn['Gioi_tinh']=='Nữ','GPA'], bins=14, alpha=.65, label='Nữ')
plt.hist(cn.loc[cn['Gioi_tinh']=='Nam','GPA'], bins=14, alpha=.65, label='Nam')
plt.axvline(cn['GPA'].mean(), linestyle='--', linewidth=1.5, label=f"TB chung = {cn['GPA'].mean():.2f}")
plt.xlabel('GPA (thang 4,0)'); plt.ylabel('Số sinh viên')
plt.title('Phân bố GPA bậc cử nhân theo giới tính')
plt.legend(frameon=False)
chart1=savefig('figure_1_gpa_hist.png')

# 2 Classification counts and rates
order=['Xuất sắc','Giỏi','Khá']
counts=cn['Loại tốt nghiệp'].value_counts().reindex(order).fillna(0)
plt.figure(figsize=(6.8,4.0))
bars=plt.bar(counts.index, counts.values)
plt.ylabel('Số sinh viên'); plt.title('Phân bố xếp loại tốt nghiệp bậc cử nhân')
for bar,val in zip(bars,counts.values):
    plt.text(bar.get_x()+bar.get_width()/2, val+3, f'{int(val)}\n({val/len(cn)*100:.1f}%)', ha='center', va='bottom')
plt.ylim(0,max(counts.values)*1.2)
chart2=savefig('figure_2_classification.png')

# 3 Graduation cohorts stacked
labels=['9/2025','10–11/2025','12/2025','1/2026','3/2026','4–5/2026','6/2026']
cnv=np.array([22,0,58,1,107,0,128])
thsv=np.array([4,0,3,0,3,0,3])
tsv=np.array([0,6,6,0,0,9,0])
plt.figure(figsize=(7.2,4.2))
plt.bar(labels,cnv,label='Cử nhân')
plt.bar(labels,thsv,bottom=cnv,label='Thạc sĩ')
plt.bar(labels,tsv,bottom=cnv+thsv,label='Tiến sĩ')
plt.ylabel('Số người tốt nghiệp'); plt.title('Số lượng tốt nghiệp theo đợt xét và bậc đào tạo')
plt.xticks(rotation=25,ha='right'); plt.legend(frameon=False,ncol=3)
chart3=savefig('figure_3_batches.png')

# 4 On-time rate by level
levels=['Cử nhân','Thạc sĩ','Tiến sĩ']
rates=[269/316*100,8/13*100,7/21*100]
plt.figure(figsize=(6.8,4.0))
bars=plt.bar(levels,rates)
plt.ylabel('Tỷ lệ đúng hạn (%)'); plt.title('Tỷ lệ tốt nghiệp đúng hạn theo bậc đào tạo')
plt.ylim(0,100)
for bar,val in zip(bars,rates):
    plt.text(bar.get_x()+bar.get_width()/2,val+2,f'{val:.1f}%',ha='center')
chart4=savefig('figure_4_ontime.png')

# 5 100% stacked classification by gender
ct=pd.crosstab(cn['Gioi_tinh'],cn['Loại tốt nghiệp']).reindex(index=['Nữ','Nam'],columns=order,fill_value=0)
pct=ct.div(ct.sum(axis=1),axis=0)*100
plt.figure(figsize=(6.8,4.1))
bottom=np.zeros(len(pct))
for col in order:
    vals=pct[col].values
    bars=plt.bar(pct.index,vals,bottom=bottom,label=col)
    for i,(bar,val) in enumerate(zip(bars,vals)):
        if val>=7:
            plt.text(bar.get_x()+bar.get_width()/2,bottom[i]+val/2,f'{val:.1f}%',ha='center',va='center',fontsize=9)
    bottom+=vals
plt.ylabel('Tỷ lệ trong từng giới (%)'); plt.title('Cơ cấu xếp loại tốt nghiệp theo giới tính')
plt.ylim(0,100); plt.legend(frameon=False,ncol=3,loc='upper center',bbox_to_anchor=(.5,-.12))
chart5=savefig('figure_5_gender_class.png')

# ---------- DOCX ----------
doc=Document()
style=doc.styles['Normal']
style.font.name='Times New Roman'; style.font.size=Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
for sec in doc.sections:
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
    sec.left_margin=Cm(3); sec.right_margin=Cm(2)

def add_p(text='', bold=False, italic=False, align=None, size=13, after=6, keep=False):
    par=doc.add_paragraph()
    run=par.add_run(text)
    run.bold=bold; run.italic=italic; run.font.name='Times New Roman'; run.font.size=Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
    if align=='center': par.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif align=='justify': par.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align=='right': par.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    par.paragraph_format.space_after=Pt(after); par.paragraph_format.line_spacing=1.15
    par.paragraph_format.keep_with_next=keep
    return par

def shade(cell, fill='D9E2F3'):
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'),fill)))

def set_repeat_table_header(row):
    row._tr.get_or_add_trPr().append(parse_xml(r'<w:tblHeader {} w:val="true"/>'.format(nsdecls('w'))))

def make_table(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(t.rows[0])
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=''; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(h)); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
    for i,row in enumerate(rows,1):
        for j,val in enumerate(row):
            c=t.rows[i].cells[j]; c.text=''; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.name='Times New Roman'; r.font.size=Pt(10)
            r._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def note(text): add_p(text,italic=True,size=10.5,after=8)

def add_figure(path,caption,width=15.5):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(3)
    p.add_run().add_picture(path,width=Cm(width))
    add_p(caption,italic=True,align='center',size=10.5,after=8)

add_p('BÁO CÁO PHÂN TÍCH KẾT QUẢ TỐT NGHIỆP NĂM HỌC 2025–2026',bold=True,align='center',size=15,after=2)
add_p('Khoa Kinh tế và Kinh doanh Quốc tế — Trường Đại học Kinh tế, ĐHQGHN',align='center',size=13,after=2)
add_p('Người thực hiện: Nguyễn Quốc Khánh | Công cụ: Python (pandas, scipy, matplotlib)',italic=True,align='center',size=11,after=14)

add_p('1. Dữ liệu và phương pháp',bold=True,size=13,keep=True)
add_p('Báo cáo sử dụng danh sách tốt nghiệp năm học 2025–2026, gồm 316 cử nhân sau khi loại 2 quan sát thiếu GPA, 13 học viên cao học và 21 nghiên cứu sinh. Dữ liệu được làm sạch và phân tích bằng Python. Các kỹ thuật được sử dụng gồm thống kê mô tả, trực quan hóa dữ liệu, kiểm định t hai mẫu độc lập theo phương pháp Welch và kiểm định Chi bình phương. Kết quả được trình bày ở dạng tổng hợp, không công bố thông tin định danh cá nhân.',align='justify')

add_p('2. Kết quả phân tích',bold=True,size=13,keep=True)
add_p('Bảng 1. Thống kê mô tả GPA bậc cử nhân (thang 4,0)',bold=True,size=12,after=4,keep=True)
make_table(['Nhóm','n','Trung bình','Độ lệch chuẩn','Min','Q1','Trung vị','Q3','Max'],[
['Toàn bộ',316,'3,360','0,291','2,51','3,23','3,40','3,57','3,88'],['Nữ',214,'3,425','0,258','2,52','3,27','3,45','3,61','3,85'],['Nam',102,'3,223','0,310','2,51','2,98','3,30','3,43','3,88']])
note('Nguồn: Tính toán từ dữ liệu tốt nghiệp bằng Python.')
add_figure(chart1,'Hình 1. Phân bố GPA bậc cử nhân theo giới tính. Nguồn: Tính toán và trực quan hóa bằng Python.')
add_p('GPA trung bình toàn bộ nhóm cử nhân đạt 3,36/4,0, trong khi trung vị là 3,40. Biểu đồ cho thấy phân phối GPA tập trung chủ yếu quanh khoảng 3,2–3,7. GPA trung bình của nữ cao hơn nam khoảng 0,20 điểm. Kiểm định Welch cho thấy chênh lệch này có ý nghĩa thống kê rất mạnh (t = 5,72; p < 0,001). Nhóm nam đồng thời có mức độ phân tán GPA lớn hơn nhóm nữ.',align='justify')

add_p('Bảng 2. Phân bố xếp loại tốt nghiệp bậc cử nhân',bold=True,size=12,after=4,keep=True)
make_table(['Xếp loại','Số lượng','Tỷ lệ (%)','GPA trung bình','GPA min–max'],[
['Xuất sắc',71,'22,5','3,691','2,89–3,88'],['Giỏi',178,'56,3','3,393','3,20–3,59'],['Khá',67,'21,2','2,922','2,51–3,58'],['Tổng',316,'100,0','3,360','2,51–3,88']])
note('Nguồn: Tính toán từ dữ liệu tốt nghiệp bằng Python.')
add_figure(chart2,'Hình 2. Phân bố xếp loại tốt nghiệp bậc cử nhân. Nguồn: Tính toán và trực quan hóa bằng Python.')
add_p('Có 78,8% sinh viên đạt loại Giỏi trở lên; trong đó nhóm Giỏi chiếm tỷ trọng lớn nhất với 56,3%. Tuy nhiên, khoảng GPA giữa các nhóm xếp loại có sự chồng lấn. Hiện tượng này có thể phản ánh các tiêu chí xếp loại ngoài GPA hoặc sai lệch nhập liệu; vì vậy, dữ liệu gốc cần được đối chiếu với quy định xếp loại và hồ sơ quyết định tốt nghiệp.',align='justify')

add_p('Bảng 3. Số lượng tốt nghiệp theo đợt xét và bậc đào tạo',bold=True,size=12,after=4,keep=True)
make_table(['Đợt xét','Cử nhân','Thạc sĩ','Tiến sĩ','Tổng','Tỷ trọng (%)'],[
['9/2025',22,4,0,26,'7,4'],['10–11/2025',0,0,6,6,'1,7'],['12/2025',58,3,6,67,'19,1'],['1/2026',1,0,0,1,'0,3'],['3/2026',107,3,0,110,'31,4'],['4–5/2026',0,0,9,9,'2,6'],['6/2026',128,3,0,131,'37,4'],['Tổng',316,13,21,350,'100,0']])
note('Nguồn: Tính toán từ dữ liệu tốt nghiệp bằng Python.')
add_figure(chart3,'Hình 3. Số lượng tốt nghiệp theo đợt xét và bậc đào tạo. Nguồn: Tính toán và trực quan hóa bằng Python.')
add_p('Hai đợt 3/2026 và 6/2026 chiếm 68,8% tổng số người tốt nghiệp. Trực quan cho thấy phần lớn quy mô ở các đợt cao điểm đến từ bậc cử nhân, trong khi tiến sĩ tập trung vào các đợt riêng từ tháng 10/2025 đến tháng 5/2026. Khoa có thể dựa trên tính mùa vụ này để bố trí nhân sự xử lý hồ sơ và tổ chức lễ tốt nghiệp.',align='justify')

add_p('Bảng 4. Tình trạng tốt nghiệp đúng hạn theo bậc đào tạo',bold=True,size=12,after=4,keep=True)
make_table(['Bậc đào tạo','Đúng hạn','Trễ hạn','Tổng','Tỷ lệ đúng hạn (%)'],[['Cử nhân*',269,47,316,'85,1'],['Thạc sĩ',8,5,13,'61,5'],['Tiến sĩ',7,14,21,'33,3']])
note('Nguồn: Tính toán bằng Python. (*) Cử nhân đúng hạn được quy ước là sinh viên khóa QH-2022 tốt nghiệp trong năm học 2025–2026; các khóa trước được xếp là trễ hạn.')
add_figure(chart4,'Hình 4. Tỷ lệ tốt nghiệp đúng hạn theo bậc đào tạo. Nguồn: Tính toán và trực quan hóa bằng Python.')
add_p('Tỷ lệ đúng hạn giảm theo bậc đào tạo: 85,1% ở cử nhân, 61,5% ở thạc sĩ và 33,3% ở tiến sĩ. Đối với bậc sau đại học, kết quả cần được diễn giải thận trọng vì tiến độ còn phụ thuộc vào đặc thù nghiên cứu, yêu cầu công bố và tình trạng vừa học vừa làm.',align='justify')

add_p('Bảng 5. Phân bố xếp loại theo giới tính và kiểm định độc lập',bold=True,size=12,after=4,keep=True)
make_table(['Giới tính','Xuất sắc','Giỏi','Khá','Tổng','Giỏi trở lên (%)'],[['Nữ','61 (28,5%)','125 (58,4%)','28 (13,1%)',214,'86,9'],['Nam','10 (9,8%)','53 (52,0%)','39 (38,2%)',102,'61,8']])
note('Nguồn: Tính toán bằng Python. Kiểm định Chi bình phương: χ²(2) = 31,87; p < 0,001.')
add_figure(chart5,'Hình 5. Cơ cấu xếp loại tốt nghiệp theo giới tính. Nguồn: Tính toán và trực quan hóa bằng Python.')
add_p('Cơ cấu xếp loại khác biệt có ý nghĩa thống kê giữa nam và nữ. Tỷ lệ đạt Giỏi trở lên ở nữ là 86,9%, cao hơn mức 61,8% ở nam. Hình 5 cho thấy khác biệt chủ yếu đến từ tỷ lệ Xuất sắc cao hơn ở nữ và tỷ lệ Khá cao hơn ở nam. Tuy nhiên, kết quả chỉ phản ánh mối liên hệ thống kê trong mẫu, không phải quan hệ nhân quả.',align='justify')

add_p('3. Kết luận và khuyến nghị',bold=True,size=13,keep=True)
add_p('Kết quả cho thấy chất lượng đầu ra bậc cử nhân ở mức cao, với GPA trung bình 3,36 và 78,8% sinh viên đạt Giỏi hoặc Xuất sắc. Hoạt động xét tốt nghiệp có tính mùa vụ rõ rệt, tập trung vào tháng 3 và tháng 6 năm 2026. Tỷ lệ tốt nghiệp đúng hạn giảm đáng kể ở các bậc đào tạo cao hơn. Bên cạnh đó, một số trường hợp GPA không tương thích trực quan với xếp loại cho thấy cần rà soát quy tắc phân loại và chuẩn hóa dữ liệu. Khoa nên xây dựng bộ quy tắc kiểm tra tự động đối với GPA, xếp loại, khóa học và tình trạng tốt nghiệp trước khi tổng hợp báo cáo chính thức.',align='justify')

doc.save(OUT)
print(OUT)
